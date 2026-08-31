# -*- coding: utf-8 -*-
"""
Gera o relatorio de implementacao em .docx.

    py tools/gerar_relatorio.py              # usa a data de hoje
    py tools/gerar_relatorio.py 2026-09-15   # usa a data informada

Sai em docs/SearchPlus_Implementacoes_AAAA-MM-DD.docx e o .pdf ao lado.
Cada rodada gera um
arquivo novo, com data no nome, em vez de sobrescrever: a ideia e ter o
historico do que foi entregue em cada momento, e nao so a foto de agora.

Capa com as cores do app; miolo todo preto. As tabelas seguem sempre o mesmo
formato -- o que mudou, como era, como e agora -- porque o documento existe
para responder "o que o usuario ganhou", e nao para repetir a especificacao,
que ja esta em docs/09-requisitos-funcionais.md.

AO ACRESCENTAR UMA FEATURE: some um capitulo em CAPITULOS, atualize os
numeros do Resumo e, se for o caso, a tabela de decisoes em aberto. O texto
de cada linha deve caber na frase "antes o usuario fazia X, agora faz Y" --
se nao couber, provavelmente e detalhe tecnico e o lugar dele e o markdown
dos requisitos.
"""
import datetime
import io
import os
import re
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# --- cores (so a capa usa) -------------------------------------------------
FUNDO = "0B0F19"
ROXO = "AB5AF7"
MAGENTA = "E879F9"
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
CINZA_CLARO = RGBColor(0x94, 0xA3, 0xB8)
PRETO = RGBColor(0, 0, 0)

HOJE = datetime.date.today()   # sobrescrito pelo argumento da linha de comando
MESES = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho",
         "agosto", "setembro", "outubro", "novembro", "dezembro"]


def sombrear(celula, hexcor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcor)
    celula._tc.get_or_add_tcPr().append(shd)


def sem_margem_interna(tabela):
    tblPr = tabela._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for lado in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:%s" % lado)
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def sem_bordas(tabela):
    tblPr = tabela._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:%s" % lado)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        b.append(el)
    tblPr.append(b)


def nao_partir(linha):
    """
    Impede a linha de quebrar entre duas paginas. Sem isso, uma linha alta
    parte no meio e a continuacao aparece com as duas primeiras colunas
    vazias -- o leitor perde a referencia do que esta lendo.
    """
    trPr = linha._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    trPr.append(el)


def repetir_cabecalho(linha):
    trPr = linha._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def interpolar(c1, c2, t):
    a = [int(c1[i:i + 2], 16) for i in (0, 2, 4)]
    b = [int(c2[i:i + 2], 16) for i in (0, 2, 4)]
    return "".join("%02X" % round(a[i] + (b[i] - a[i]) * t) for i in range(3))


# ===========================================================================
# CONTEUDO
# ===========================================================================
# (titulo, paragrafo de contexto, [(o que, antes, agora)], "faixa de RFs")
CAPITULOS = [
    ("1. Encontrar as coisas: a busca",
     "A busca é o coração do produto. As mudanças aqui foram menos sobre achar "
     "melhor e mais sobre não obrigar o usuário a recomeçar do zero quando o "
     "resultado não é o esperado.",
     [
         ("Limpar o campo",
          "Apagava letra por letra, ou selecionava tudo e deletava.",
          "Um × dentro do campo limpa de uma vez e deixa o cursor pronto para digitar. "
          "Os resultados na tela continuam lá: limpar o texto e sair da tela de "
          "resultados são coisas diferentes."),
         ("Tirar um assunto do resultado",
          "Se a busca trazia coisa demais, era preciso reescrever a frase inteira.",
          "Escrever “-palavra” tira aquele assunto. Vale para o texto e também para a "
          "imagem, porque a maioria das fotos com gente não tem a palavra “pessoas” "
          "escrita em lugar nenhum."),
         ("Pesquisar dentro do que já apareceu",
          "Cada nova busca varria a biblioteca inteira outra vez.",
          "“Buscar dentro destes” limita a próxima busca ao que está na tela — útil "
          "para ir afunilando sem perder o caminho."),
         ("Ver e desfazer os filtros",
          "Não dava para saber o que estava filtrando o resultado.",
          "Uma trilha mostra cada refinamento aplicado, e cada um sai sozinho com um "
          "clique. Aparece também quando a busca não achou nada, que é justamente "
          "quando o caminho de volta é remover um filtro."),
         ("Buscar só entre os favoritos",
          "O favorito só servia para reencontrar depois, na aba de favoritos.",
          "Um botão “Favoritos” na barra restringe a busca ao que você marcou, e "
          "combina com os outros filtros em vez de substituí-los."),
         ("Saber por que um resultado apareceu",
          "O resultado aparecia sem explicação.",
          "Cada um traz uma etiqueta com o motivo: a aparência, o que a imagem mostra, "
          "o texto do documento ou o nome do arquivo."),
         ("Quando não acha nada",
          "A tela dizia apenas “nada encontrado”.",
          "Explica o motivo provável e sugere categorias que têm conteúdo, além das "
          "buscas recentes — caminhos que levam a resultado de verdade."),
     ],
     "RF-001 a RF-012 · RF-193 a RF-199 · RF-219 a RF-233 · RF-265 a RF-267"),

    ("2. A tela inicial",
     "Muita gente não usa a busca escrita: navega pelas categorias que a análise "
     "monta sozinha. A tela inicial passou a ser um lugar de trabalho, e não só "
     "uma vitrine.",
     [
         ("Ver a home logo ao entrar",
          "As pastas importadas só apareciam depois da primeira busca.",
          "As categorias aparecem desde o login."),
         ("Escolher quais pastas aparecem",
          "Tudo o que foi importado se misturava nas mesmas categorias, sem dizer de "
          "onde cada imagem veio.",
          "Um botão no canto lista as pastas importadas com nome e quantidade, e você "
          "escolhe quais mostrar. A escolha fica guardada na conta, não no navegador."),
         ("Deixar a home limpa",
          "Desmarcar todas as pastas voltava a mostrar tudo — a escolha era ignorada.",
          "Nenhuma pasta marcada mostra nenhuma pasta, para quem quer só a barra de "
          "busca. Há um atalho “Não mostrar nenhuma” no menu."),
         ("Selecionar sem passar pela busca",
          "Para juntar fotos numa coleção era preciso buscar antes.",
          "Cada card tem caixa de seleção, e cada categoria tem “Selecionar tudo”. A "
          "mesma foto que aparece em duas categorias marca nas duas."),
         ("Favoritar na home",
          "Era preciso abrir a imagem para marcar a estrela.",
          "O card tem botão de favoritar, e marcar atualiza todos os cards daquela "
          "imagem de uma vez."),
     ],
     "RF-234 a RF-249 · RF-259 a RF-264"),

    ("3. Selecionar imagens",
     "Selecionar e favoritar são coisas diferentes: uma é um passo de trabalho, "
     "a outra é uma marcação que dura. O app passou a tratá-las assim.",
     [
         ("Selecionar sem favoritar",
          "Só existia o coração, que serve para outra finalidade.",
          "Caixa de seleção própria, à esquerda do card. Selecionar não favorita e "
          "desfavoritar não tira da seleção."),
         ("Marcar tudo de uma vez",
          "Marcava-se uma a uma.",
          "Um botão marca todos os resultados da busca atual; o mesmo botão desmarca "
          "quando tudo já está marcado. O rótulo diz o que o clique vai fazer."),
         ("Não perder a seleção",
          "Trocar de filtro ou recarregar a página apagava tudo.",
          "A seleção sobrevive à troca de filtro e ao recarregamento. Fechar a aba "
          "limpa: continua sendo passo de trabalho, não algo para guardar."),
         ("Saber quantas estão marcadas",
          "Não havia contagem.",
          "A barra mostra quantos resultados existem e quantos estão marcados."),
         ("Arquivo que sumiu do disco",
          "Entraria na coleção como um item quebrado.",
          "Ao restaurar a seleção, o que não existe mais é descartado e você é avisado "
          "de quantos saíram — para não estranhar o número menor."),
     ],
     "RF-013 a RF-021 · RF-065 a RF-072 · RF-278 a RF-283"),

    ("4. Coleções",
     "A coleção é o agrupamento manual, feito por você, ao lado das categorias "
     "automáticas. Ganhou as operações que faltavam para ser usável no dia a dia.",
     [
         ("Adicionar várias de uma vez",
          "Uma imagem por vez, abrindo cada uma.",
          "Tudo o que está selecionado entra numa ação só, e o app informa quantas "
          "entraram e quantas já estavam lá."),
         ("Renomear",
          "Não era possível.",
          "Renomeia pelo cabeçalho da coleção ou pelas configurações. As pastas já "
          "exportadas acompanham o novo nome."),
         ("Ordenar a lista",
          "Ordem fixa.",
          "Por mais recentes, mais antigas, nome ou quantidade de itens. A escolha "
          "sobrevive ao recarregamento."),
         ("Escolher a capa",
          "A capa era sempre o mosaico automático das últimas imagens.",
          "Dá para escolher a imagem da capa. O mosaico continua sendo o padrão, e a "
          "coleção volta a ele sozinha se a imagem escolhida sair da biblioteca."),
         ("Transformar favoritos em coleção",
          "Os favoritos não conversavam com as coleções.",
          "Na aba de favoritos, selecione um ou todos e mande para uma coleção "
          "existente, ou crie uma na hora. Ao terminar, o app leva você até ela."),
         ("Abrir a lista de coleções",
          "A tela fazia uma consulta ao banco por coleção; com muitas, demorava.",
          "Uma consulta só, independente de quantas coleções existam."),
     ],
     "RF-022 a RF-032 · RF-151 a RF-153 · RF-250 a RF-258 · RF-268 a RF-276"),

    ("5. Exportar para uma pasta do computador",
     "Exportar é copiar as imagens da coleção para uma pasta comum do Windows, "
     "para levar num pendrive, mandar para alguém ou abrir em outro programa. "
     "Os arquivos originais nunca são movidos nem apagados.",
     [
         ("Exportar logo depois de adicionar",
          "Adicionar e exportar eram passos separados; era fácil esquecer o segundo.",
          "Terminada a adição, o app oferece exportar aquela coleção na hora."),
         ("Escolher o que sai",
          "Saía tudo, do jeito que estava.",
          "Só imagens, só documentos ou tudo. Dá para renomear em lote com um padrão, "
          "reduzir o tamanho das imagens e organizar a saída em subpastas por mês."),
         ("Nome dos arquivos",
          "O nome original, e ponto.",
          "Padrão configurável, com numeração que usa zeros à esquerda — sem eles o "
          "Explorer ordena 1, 10, 11, 2. A extensão nunca vem do padrão: trocá-la não "
          "converte o arquivo, só faz abrir com o programa errado."),
         ("Escolher a pasta de destino",
          "Escolhia-se a pasta a cada exportação.",
          "A última pasta usada já vem preenchida."),
         ("Exportar a mesma coleção de novo",
          "Criava outra pasta sem avisar, e não dava para saber qual era qual.",
          "O app avisa que a coleção já tem pasta e quantas. O nome da coleção é sempre "
          "o começo do nome da pasta, e você escolhe o complemento em vez da "
          "numeração automática."),
         ("Saber o que já foi exportado",
          "Exportou e acabou; não ficava registro.",
          "Um histórico mostra o que foi, quando, para onde e o que falhou — e diz se "
          "a pasta ainda existe, para não oferecer um botão que não abre nada."),
         ("Consertar o que falhou",
          "Era preciso refazer a exportação inteira.",
          "“Tentar de novo” repete só os que falharam, na mesma pasta. “Tirar da "
          "coleção os que sumiram” limpa os arquivos que não estão mais no disco — "
          "conferindo na hora, porque o disco externo pode ter sido reconectado."),
     ],
     "RF-033 a RF-064 · RF-073 a RF-079 · RF-284 a RF-306"),

    ("6. Pasta vinculada: a coleção que se mantém sozinha",
     "Este foi o incômodo mais concreto relatado: adicionar uma foto nova a uma "
     "coleção já exportada não colocava a foto na pasta. A coleção e a pasta "
     "passaram a andar juntas.",
     [
         ("Quando o app pergunta o destino",
          "Perguntava a cada foto adicionada — repetitivo e fácil de errar.",
          "Pergunta uma vez, na criação da coleção, com três opções: enviar sempre sem "
          "perguntar, perguntar antes de cada envio, ou não enviar. Uma animação curta "
          "mostra o que cada opção faz."),
         ("Mudar de ideia depois",
          "A escolha ficava presa ao momento da criação.",
          "A aba “Configurações” dentro da coleção muda o modo e o destino a qualquer "
          "momento, e mostra qual está em vigor."),
         ("Foto nova numa coleção já exportada",
          "A foto entrava na coleção mas não aparecia na pasta.",
          "No modo “enviar sempre”, é copiada na hora — e só ela, não a coleção "
          "inteira. Adicionar 1 foto a uma coleção de 300 copia 1 arquivo."),
         ("Mais de uma pasta de destino",
          "Um destino só por coleção.",
          "Dá para marcar várias pastas para receber ao mesmo tempo, espelhando a "
          "coleção em mais de um lugar — ou nenhuma, parando só o envio automático "
          "sem perder as pastas já criadas."),
         ("Remover uma imagem da coleção",
          "A cópia continuava na pasta exportada, e a pasta ia ficando desatualizada.",
          "A cópia sai junto, conforme o modo escolhido. O arquivo original nas suas "
          "pastas nunca é tocado."),
         ("Ver o que já foi para a pasta",
          "Não havia como conferir sem abrir o Explorer e comparar na mão.",
          "“Status da pasta” mostra, por pasta, o que já foi copiado e o que falta, com "
          "nomes — e copia as faltantes ali mesmo. Também lista arquivos que estão na "
          "pasta mas não na coleção."),
         ("Excluir a coleção",
          "A coleção sumia e as pastas ficavam no disco sem explicação.",
          "O app pergunta o que fazer com cada pasta, mostrando caminho completo e "
          "quantidade de arquivos. Manter é uma opção de primeira classe. Apagar exige "
          "duas etapas, e o aviso diz que não vai para a Lixeira do Windows."),
     ],
     "RF-080 a RF-146"),

    ("7. A análise das imagens",
     "Ao importar uma pasta, cada arquivo passa por uma análise que é o que "
     "permite buscar por significado depois. É demorada, e o problema era não "
     "saber o que estava acontecendo.",
     [
         ("Saber quanto falta",
          "A barra dizia “N na fila”, que não responde a pergunta real: dá tempo de "
          "almoçar?",
          "Mostra o tempo restante, calculado pelo ritmo que a análise está de fato "
          "conseguindo na sua máquina. Abaixo de um minuto, diz “quase terminando” em "
          "vez de um número."),
         ("Saber o que entrou e o que ficou de fora",
          "A análise terminava e pronto.",
          "Um resumo por pasta com indexados, ignorados e com erro — separados, porque "
          "um .zip no meio das fotos não é defeito. Os erros abrem com nome e motivo. "
          "Fica guardado e reabre pelas Configurações."),
         ("Pasta que mudou depois de importada",
          "Só reimportando a pasta inteira.",
          "Um botão “Verificar” em cada pasta acha o que é novo, o que foi alterado e "
          "o que sumiu, e responde na hora com os números e os nomes."),
         ("Arquivo que sumiu do disco",
          "Não havia tratamento.",
          "É marcado, nunca apagado: o motivo mais comum é um disco externo "
          "desconectado, e apagar jogaria fora a descrição e a participação em "
          "coleções. Volta sozinho quando reaparece."),
     ],
     "RF-154 a RF-168 · RF-200 a RF-218"),

    ("8. Abrir o programa",
     "O aplicativo carrega modelos de análise pesados ao iniciar. Antes, ele "
     "carregava tudo antes de responder qualquer coisa.",
     [
         ("Esperar para ver a tela",
          "Cerca de 30 segundos sem nada. O navegador dizia “não foi possível acessar "
          "este site” e a conclusão natural era que o programa não tinha aberto.",
          "A tela aparece em cerca de 2 segundos, e a análise continua carregando por "
          "trás. Medido: 29,7s para ~2,2s."),
         ("Buscar antes de tudo carregar",
          "A busca podia devolver uma lista vazia, que se lê como “não tem nada”.",
          "Um aviso discreto diz que a busca está sendo preparada e some sozinho. "
          "Buscar cedo demais recebe “tente de novo em instantes”."),
         ("Mensagens de erro",
          "Citavam nomes internos, como “SBERT indisponível”.",
          "Nenhuma mensagem cita nome de modelo ou termo técnico."),
     ],
     "RF-169 a RF-179"),

    ("9. Desfazer",
     "Excluir era definitivo e imediato. Passou a haver uma rede de proteção em "
     "duas camadas.",
     [
         ("Excluir uma coleção por engano",
          "Excluiu, perdeu.",
          "Um aviso com “desfazer” fica 8 segundos na tela. Quem perder esse tempo "
          "encontra o item na lixeira, dentro das Configurações, por 30 dias."),
         ("Remover imagens de uma coleção",
          "Mesma situação.",
          "Mesmo desfazer."),
         ("O que volta ao restaurar",
          "—",
          "A coleção volta com o mesmo identificador, com as pastas geradas e com o "
          "modo de envio. Imagem que você apagou da biblioteca nesse meio-tempo é "
          "pulada, em vez de fazer a restauração inteira falhar."),
         ("Descartar da lixeira",
          "—",
          "É a única exclusão sem desfazer do app, e pede confirmação. É assim de "
          "propósito: a lixeira é o desfazer."),
     ],
     "RF-180 a RF-192"),

    ("10. Acessibilidade",
     "Um conjunto de correções para que o app funcione para quem usa leitor de "
     "tela, navega por teclado ou não distingue certas cores.",
     [
         ("Escolher uma cor ilegível",
          "Era possível escolher qualquer cor de destaque, inclusive uma que não dava "
          "para ler.",
          "O contraste é medido no momento da escolha, contra o fundo em que a cor vai "
          "ser lida de verdade. Abaixo do mínimo, o app avisa e oferece o tom mais "
          "próximo que passa — quem escolheu aquela cor queria aquela cor."),
         ("O tema que já vinha instalado",
          "Três combinações do tema padrão estavam abaixo do mínimo de leitura.",
          "Auditado e corrigido."),
         ("Mudanças na tela",
          "Passavam em silêncio para quem usa leitor de tela.",
          "Progresso da análise, contagem de resultados e avisos são anunciados. Erro "
          "interrompe a leitura; o resto espera a pausa."),
         ("Teclado nas janelas",
          "Das 23 janelas do app, 14 fechavam com Esc; nove não tinham tecla nenhuma.",
          "Todas fecham com Esc. O foco entra na janela ao abrir, fica preso dentro "
          "dela enquanto está aberta e volta ao botão que a abriu."),
         ("Descrição das imagens",
          "As imagens iam sem texto alternativo.",
          "A descrição gerada pela análise vira o texto alternativo."),
         ("Favorito sem indicação",
          "Favoritava e o botão ficava vazio — um círculo em branco, sem informação.",
          "O coração passa de vazado a preenchido: a diferença não é só de cor, então "
          "funciona para quem não distingue as duas."),
     ],
     "RF-147 a RF-150 · RF-307 a RF-327 · RF-335"),

    ("11. Aparência dos ícones",
     "Os ícones da interface eram emoji. Foram substituídos por um conjunto "
     "desenhado, com 35 símbolos no mesmo traço.",
     [
         ("Ícone e tema",
          "Emoji são coloridos pela fonte do sistema. Ignoravam o tema e a cor de "
          "destaque escolhida, e não dava para corrigir o contraste deles como "
          "corrigimos o do texto.",
          "Cada ícone acompanha a cor do texto ao lado, em qualquer tema."),
         ("Ícone e máquina",
          "O desenho mudava entre Windows, Android e cada navegador. Em fonte antiga, "
          "alguns viravam um retângulo vazio.",
          "O desenho é o mesmo em qualquer lugar, porque vai junto com a página."),
         ("Botões que sumiam",
          "Dois botões ficavam pretos sobre fundo escuro — praticamente invisíveis. O "
          "emoji escondia o problema porque trazia a própria cor.",
          "Corrigido, com verificação automática de contraste."),
     ],
     "RF-328 a RF-335"),

    ("12. Segurança e comportamento interno",
     "Mudanças que o usuário não vê diretamente, mas que protegem os arquivos e "
     "a conta.",
     [
         ("Abrir um arquivo pelo app",
          "O endereço aceitava qualquer caminho do computador.",
          "Só abre arquivos que estejam dentro das pastas que você importou."),
         ("Chave de sessão",
          "Era um valor fixo escrito no código.",
          "Gerada e guardada na primeira vez que o programa roda."),
         ("Atualizações da interface",
          "O navegador guardava a versão antiga em cache e a mudança não aparecia.",
          "Os arquivos da interface não são mais guardados em cache."),
         ("Apagar pastas",
          "—",
          "O servidor recusa apagar qualquer caminho que não esteja registrado para "
          "aquela coleção e aquele usuário, mesmo que exista no disco."),
     ],
     "RF-045 · RF-100 · RF-107 a RF-108 · RF-138 a RF-139"),
]


def converter_para_pdf(caminho_docx):
    """
    Gera o PDF ao lado do .docx.

    Duas rotas, nesta ordem, porque nem toda maquina tem as duas:

      1. Word, se estiver instalado. E o que produz o resultado fiel -- foi com
         ele que a paginacao e a capa foram conferidas.
      2. LibreOffice, como alternativa. Pagina de forma ligeiramente diferente
         (a contagem de paginas pode variar), mas resolve numa maquina sem Word.

    Sem nenhum dos dois, avisa e devolve None em vez de derrubar: o .docx, que
    e a entrega, ja esta gravado a essa altura.
    """
    absoluto = os.path.abspath(caminho_docx)
    destino = absoluto[:-5] + ".pdf"

    # --- 1. Word ----------------------------------------------------------
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            # ReadOnly=True para nao mexer no arquivo que acabou de sair.
            documento = word.Documents.Open(absoluto, False, True)
            documento.SaveAs(destino, FileFormat=17)   # 17 = PDF
            documento.Close(False)
        finally:
            word.Quit()
        return destino
    except Exception as erro:
        print("  Word nao converteu (%s); tentando o LibreOffice." % erro)

    # --- 2. LibreOffice ---------------------------------------------------
    for candidato in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "soffice",
    ):
        try:
            subprocess.run(
                [candidato, "--headless", "--convert-to", "pdf",
                 "--outdir", os.path.dirname(absoluto), absoluto],
                check=True, capture_output=True, timeout=180)
            if os.path.exists(destino):
                return destino
        except (OSError, subprocess.SubprocessError):
            continue

    print("  Nenhum conversor disponivel; so o .docx foi gerado.")
    return None


def montar(hoje=None):
    global HOJE
    if hoje is not None:
        HOJE = hoje

    doc = Document()

    # --- estilos globais: tudo preto -------------------------------------
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = PRETO
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for nome, tam in (("Heading 1", 16), ("Heading 2", 12.5)):
        est = doc.styles[nome]
        est.font.name = "Calibri"
        est.font.size = Pt(tam)
        est.font.bold = True
        est.font.color.rgb = PRETO
        est.paragraph_format.space_before = Pt(18)
        est.paragraph_format.space_after = Pt(6)

    # =====================================================================
    # CAPA
    # =====================================================================
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    for lado in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, lado, Cm(0))

    capa = doc.add_table(rows=1, cols=1)
    sem_bordas(capa)
    sem_margem_interna(capa)
    cel = capa.cell(0, 0)
    cel.width = Cm(21)
    capa.rows[0].height = Cm(29.4)
    cel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    sombrear(cel, FUNDO)

    def linha_capa(texto, tamanho, cor, espaco_antes=0, espaco_depois=0,
                   negrito=False, maiuscula_espacada=False):
        p = cel.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(espaco_antes)
        p.paragraph_format.space_after = Pt(espaco_depois)
        r = p.add_run(texto)
        r.font.size = Pt(tamanho)
        r.font.color.rgb = cor
        r.font.bold = negrito
        r.font.name = "Calibri"
        if maiuscula_espacada:
            rPr = r._element.get_or_add_rPr()
            sp = OxmlElement("w:spacing")
            sp.set(qn("w:val"), "60")
            rPr.append(sp)
        return p

    # o paragrafo que a celula ja tem, vazio, serve de respiro no topo
    cel.paragraphs[0].paragraph_format.space_after = Pt(0)

    p = cel.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SEARCH")
    r.font.size, r.font.bold, r.font.name = Pt(46), True, "Calibri"
    r.font.color.rgb = BRANCO
    r2 = p.add_run("+")
    r2.font.size, r2.font.bold, r2.font.name = Pt(46), True, "Calibri"
    r2.font.color.rgb = RGBColor.from_string(ROXO)

    # faixa com o degrade do tema
    faixa_wrap = cel.add_paragraph()
    faixa_wrap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    faixa_wrap.paragraph_format.space_before = Pt(14)
    faixa_wrap.paragraph_format.space_after = Pt(0)

    faixa = cel.add_table(rows=1, cols=12)
    sem_bordas(faixa)
    sem_margem_interna(faixa)
    faixa.rows[0].height = Cm(0.18)
    faixa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(12):
        c = faixa.cell(0, i)
        c.width = Cm(9.0 / 12)
        sombrear(c, interpolar(ROXO, MAGENTA, i / 11.0))
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        run = c.paragraphs[0].add_run(" ")
        run.font.size = Pt(3)

    linha_capa("RELATÓRIO DE IMPLEMENTAÇÃO", 11, RGBColor.from_string(MAGENTA),
               espaco_antes=26, espaco_depois=4, negrito=True,
               maiuscula_espacada=True)
    linha_capa("Requisitos funcionais entregues", 20, BRANCO,
               espaco_antes=6, espaco_depois=2)
    linha_capa("e o que mudou para quem usa", 20, BRANCO, espaco_depois=0)

    linha_capa("%d de %s de %d" % (HOJE.day, MESES[HOJE.month - 1], HOJE.year),
               12, CINZA_CLARO, espaco_antes=40, espaco_depois=2)
    linha_capa("338 requisitos · 799 testes automatizados",
               10.5, CINZA_CLARO, espaco_depois=2)
    linha_capa("branch feature/colecoes-organizacao-e-pastas",
               9.5, CINZA_CLARO, espaco_depois=0)

    # =====================================================================
    # MIOLO
    # =====================================================================
    corpo = doc.add_section(WD_SECTION.NEW_PAGE)
    corpo.page_width, corpo.page_height = Cm(21), Cm(29.7)
    corpo.top_margin = corpo.bottom_margin = Cm(2.2)
    corpo.left_margin = corpo.right_margin = Cm(2.5)

    doc.add_heading("Como ler este documento", level=1)
    for t in [
        "Este relatório reúne o que foi implementado nas três frentes de trabalho e "
        "descreve cada mudança pelo que ela significa na prática: como era antes e "
        "como ficou. A linguagem é deliberadamente simples, sem termos técnicos.",
        "Cada capítulo termina com os códigos dos requisitos que ele cobre, para quem "
        "quiser rastrear até a especificação. O detalhamento técnico de cada requisito "
        "continua em docs/09-requisitos-funcionais.md, no repositório — este documento "
        "não o repete de propósito, porque 338 linhas de especificação não se leem "
        "como um relato.",
        "As linhas marcadas com “—” na coluna “Como era” são recursos que simplesmente "
        "não existiam antes.",
    ]:
        doc.add_paragraph(t)

    doc.add_heading("Resumo", level=1)
    resumo = doc.add_table(rows=1, cols=2)
    resumo.style = "Table Grid"
    resumo.autofit = False
    cab = resumo.rows[0]
    repetir_cabecalho(cab)
    for i, txt in enumerate(("Indicador", "Número")):
        cab.cells[i].text = ""
        run = cab.cells[i].paragraphs[0].add_run(txt)
        run.bold = True
    for rot, val in [
        ("Requisitos funcionais especificados", "338"),
        ("Requisitos implementados ou corrigidos", "337"),
        ("Requisitos fora de escopo, por decisão registrada", "1"),
        ("Testes automatizados passando", "799"),
        ("Áreas do produto afetadas", "12"),
    ]:
        l = resumo.add_row()
        nao_partir(l)
        l.cells[0].text = rot
        l.cells[1].text = val
    for l in resumo.rows:
        l.cells[0].width = Cm(11.5)
        l.cells[1].width = Cm(4.5)

    # --- capitulos -------------------------------------------------------
    for titulo, contexto, linhas, rfs in CAPITULOS:
        doc.add_heading(titulo, level=1)
        pc = doc.add_paragraph(contexto)
        pc.paragraph_format.keep_with_next = True

        tab = doc.add_table(rows=1, cols=3)
        tab.style = "Table Grid"
        tab.autofit = False
        cab = tab.rows[0]
        repetir_cabecalho(cab)
        for i, txt in enumerate(("O que mudou", "Como era", "Como é agora")):
            cab.cells[i].text = ""
            run = cab.cells[i].paragraphs[0].add_run(txt)
            run.bold = True

        for oque, antes, agora in linhas:
            l = tab.add_row()
            nao_partir(l)
            l.cells[0].paragraphs[0].add_run(oque).bold = True
            l.cells[1].text = antes
            l.cells[2].text = agora

        for l in tab.rows:
            l.cells[0].width = Cm(3.6)
            l.cells[1].width = Cm(6.2)
            l.cells[2].width = Cm(6.2)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run("Requisitos: " + rfs)
        r.font.size = Pt(8.5)
        r.italic = True

    # --- fora de escopo --------------------------------------------------
    doc.add_heading("O que ficou de fora, e por quê", level=1)
    fora = doc.add_table(rows=1, cols=2)
    fora.style = "Table Grid"
    fora.autofit = False
    cab = fora.rows[0]
    repetir_cabecalho(cab)
    for i, txt in enumerate(("Item", "Motivo")):
        cab.cells[i].text = ""
        cab.cells[i].paragraphs[0].add_run(txt).bold = True
    for item, motivo in [
        ("Coleções dentro de coleções",
         "Fora de escopo por decisão registrada no planejamento (RF-277). Muda a "
         "estrutura de dados e a navegação inteira; é uma frente própria."),
        ("Remover arquivos duplicados na exportação",
         "Dois itens diferentes da coleção que apontem para arquivos de conteúdo "
         "idêntico são exportados os dois. Comparar conteúdo é uma decisão de produto "
         "que ainda não foi tomada."),
    ]:
        l = fora.add_row()
        nao_partir(l)
        l.cells[0].paragraphs[0].add_run(item).bold = True
        l.cells[1].text = motivo
    for l in fora.rows:
        l.cells[0].width = Cm(4.6)
        l.cells[1].width = Cm(11.4)

    doc.add_heading("Decisões que ainda dependem de você", level=1)
    doc.add_paragraph(
        "Pontos em aberto por escolha, não por esquecimento. Nenhum bloqueia o "
        "uso do app; todos mudam alguma coisa se você decidir diferente.")
    dec = doc.add_table(rows=1, cols=2)
    dec.style = "Table Grid"
    dec.autofit = False
    cab = dec.rows[0]
    repetir_cabecalho(cab)
    for i, txt in enumerate(("Questão", "Situação hoje")):
        cab.cells[i].text = ""
        cab.cells[i].paragraphs[0].add_run(txt).bold = True
    for q, sit in [
        ("A tecla Esc deveria limpar a busca?",
         "Hoje Esc fecha janelas em todo o app, e usar a mesma tecla para duas "
         "coisas confundiria. O × do campo limpa. Pode mudar depois sem retrabalho."),
        ("Cor dos botões com degradê",
         "Ficaram um tom mais escuros para o texto branco atingir o mínimo de "
         "leitura. Se achar que descaracterizou a marca, dá para reverter — mas o "
         "texto dos botões volta a ficar abaixo do mínimo. É decisão de produto."),
        ("Suporte a outros sistemas",
         "O app é de Windows: usa o Explorer e os seletores de pasta do sistema. "
         "Abrir para Mac e Linux é possível, e fica mais caro quanto mais tarde."),
        ("Testes automatizados da interface",
         "Adicionaria uma ferramenta nova ao projeto, então não fiz por conta "
         "própria. Hoje a interface é conferida manualmente e por checagens "
         "automáticas rodadas no navegador."),
    ]:
        l = dec.add_row()
        nao_partir(l)
        l.cells[0].paragraphs[0].add_run(q).bold = True
        l.cells[1].text = sit
    for l in dec.rows:
        l.cells[0].width = Cm(5.0)
        l.cells[1].width = Cm(11.0)

    doc.add_heading("Onde encontrar o detalhe", level=1)
    ref = doc.add_table(rows=1, cols=2)
    ref.style = "Table Grid"
    ref.autofit = False
    cab = ref.rows[0]
    repetir_cabecalho(cab)
    for i, txt in enumerate(("Assunto", "Arquivo no repositório")):
        cab.cells[i].text = ""
        cab.cells[i].paragraphs[0].add_run(txt).bold = True
    for a, b in [
        ("Especificação de cada requisito", "docs/09-requisitos-funcionais.md"),
        ("Desempenho, segurança e limites", "docs/10-requisitos-nao-funcionais.md"),
        ("Endereços que a interface usa", "docs/API.md"),
        ("Como substituir ou mexer na interface", "docs/FRONTEND.md"),
        ("Como rodar os testes", "docs/TESTING.md"),
    ]:
        l = ref.add_row()
        nao_partir(l)
        l.cells[0].text = a
        l.cells[1].text = b
    for l in ref.rows:
        l.cells[0].width = Cm(7.0)
        l.cells[1].width = Cm(9.0)

    nome = "docs/SearchPlus_Implementacoes_%s.docx" % HOJE.isoformat()
    doc.save(nome)
    return nome, converter_para_pdf(nome)


if __name__ == "__main__":
    import sys

    # Roda a partir da raiz do repositorio, nao de onde o usuario chamou.
    os.chdir(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

    data = datetime.date.fromisoformat(sys.argv[1]) if len(sys.argv) > 1 else None
    for caminho in montar(data):
        if caminho:
            print(os.path.relpath(caminho).replace("\\", "/"))
