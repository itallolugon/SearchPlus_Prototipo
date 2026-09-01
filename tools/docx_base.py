# -*- coding: utf-8 -*-
"""
Base comum dos documentos em Word do projeto.

Existe porque os relatorios sao periodicos: capa, estilos, tabelas e conversao
para PDF sao sempre os mesmos, e duplicar isso a cada documento novo faz os
documentos irem divergindo de aparencia sem ninguem decidir isso.

Quem escreve um documento novo cuida do CONTEUDO; a forma vem daqui.

    from docx_base import novo_documento, capa, tabela, salvar

O padrao visual, decidido uma vez: capa com as cores do app, miolo todo preto,
tabelas de grade simples sem preenchimento.
"""
import datetime
import os
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# --- cores do app (so a capa usa) ------------------------------------------
FUNDO = "0B0F19"
ROXO = "AB5AF7"
MAGENTA = "E879F9"
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
CINZA_CLARO = RGBColor(0x94, 0xA3, 0xB8)
PRETO = RGBColor(0, 0, 0)

MESES = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho",
         "agosto", "setembro", "outubro", "novembro", "dezembro"]


# ===========================================================================
# XML de baixo nivel que o python-docx nao expoe
# ===========================================================================
def sombrear(celula, hexcor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcor)
    celula._tc.get_or_add_tcPr().append(shd)


def sem_margem_interna(tabela):
    mar = OxmlElement("w:tblCellMar")
    for lado in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:%s" % lado)
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tabela._tbl.tblPr.append(mar)


def sem_bordas(tabela):
    b = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:%s" % lado)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        b.append(el)
    tabela._tbl.tblPr.append(b)


def repetir_cabecalho(linha):
    """Repete a primeira linha no alto de cada pagina que a tabela ocupa."""
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    linha._tr.get_or_add_trPr().append(el)


def nao_partir(linha):
    """
    Impede a linha de quebrar entre duas paginas. Sem isso, uma linha alta
    parte no meio e a continuacao aparece com as primeiras colunas vazias --
    o leitor perde a referencia do que esta lendo.
    """
    linha._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def _interpolar(c1, c2, t):
    a = [int(c1[i:i + 2], 16) for i in (0, 2, 4)]
    b = [int(c2[i:i + 2], 16) for i in (0, 2, 4)]
    return "".join("%02X" % round(a[i] + (b[i] - a[i]) * t) for i in range(3))


# ===========================================================================
# Documento
# ===========================================================================
def novo_documento():
    """Documento A4 com os estilos do projeto: tudo preto, Calibri."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = PRETO
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for nome, tam in (("Heading 1", 16), ("Heading 2", 12.5), ("Heading 3", 11)):
        est = doc.styles[nome]
        est.font.name = "Calibri"
        est.font.size = Pt(tam)
        est.font.bold = True
        est.font.color.rgb = PRETO
        est.paragraph_format.space_before = Pt(16)
        est.paragraph_format.space_after = Pt(6)

    return doc


def capa(doc, chapeu, titulo, subtitulo, data, rodape=()):
    """
    Primeira pagina: fundo escuro sangrando ate a borda.

    E uma tabela de uma celula, e nao cor de fundo de pagina, porque o Word so
    imprime cor de fundo se a pessoa ligar uma opcao especifica -- a capa
    sairia branca na impressao de quem nao sabe disso.
    """
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    for lado in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, lado, Cm(0))

    tab = doc.add_table(rows=1, cols=1)
    sem_bordas(tab)
    sem_margem_interna(tab)
    cel = tab.cell(0, 0)
    cel.width = Cm(21)
    tab.rows[0].height = Cm(29.4)
    cel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    sombrear(cel, FUNDO)

    def linha(texto, tamanho, cor, antes=0, depois=0, negrito=False, espacada=False):
        p = cel.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(antes)
        p.paragraph_format.space_after = Pt(depois)
        r = p.add_run(texto)
        r.font.size = Pt(tamanho)
        r.font.color.rgb = cor
        r.font.bold = negrito
        r.font.name = "Calibri"
        if espacada:
            sp = OxmlElement("w:spacing")
            sp.set(qn("w:val"), "60")
            r._element.get_or_add_rPr().append(sp)

    # marca
    p = cel.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("SEARCH")
    r.font.size, r.font.bold, r.font.name = Pt(46), True, "Calibri"
    r.font.color.rgb = BRANCO
    r2 = p.add_run("+")
    r2.font.size, r2.font.bold, r2.font.name = Pt(46), True, "Calibri"
    r2.font.color.rgb = RGBColor.from_string(ROXO)

    # faixa com o degrade do tema
    espaco = cel.add_paragraph()
    espaco.alignment = WD_ALIGN_PARAGRAPH.CENTER
    espaco.paragraph_format.space_before = Pt(14)
    espaco.paragraph_format.space_after = Pt(0)

    faixa = cel.add_table(rows=1, cols=12)
    sem_bordas(faixa)
    sem_margem_interna(faixa)
    faixa.rows[0].height = Cm(0.18)
    faixa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(12):
        c = faixa.cell(0, i)
        c.width = Cm(9.0 / 12)
        sombrear(c, _interpolar(ROXO, MAGENTA, i / 11.0))
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        c.paragraphs[0].add_run(" ").font.size = Pt(3)

    linha(chapeu, 11, RGBColor.from_string(MAGENTA),
          antes=26, depois=4, negrito=True, espacada=True)
    for i, parte in enumerate(titulo):
        linha(parte, 20, BRANCO, antes=6 if i == 0 else 0, depois=0)
    if subtitulo:
        linha(subtitulo, 11.5, CINZA_CLARO, antes=14, depois=0)

    linha("%d de %s de %d" % (data.day, MESES[data.month - 1], data.year),
          12, CINZA_CLARO, antes=36, depois=2)
    for t in rodape:
        linha(t, 9.5, CINZA_CLARO, depois=2)


def abrir_miolo(doc):
    """Secao do conteudo: margens normais, comeca em pagina nova."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.5)
    return sec


def tabela(doc, cabecalhos, linhas, larguras, negrito_primeira=True):
    """
    Tabela de grade simples, texto preto, sem preenchimento.

    `linhas` e uma lista de tuplas de texto. Uma celula que comece com "**"
    sai em negrito -- serve para a primeira coluna, que costuma ser o rotulo.
    """
    tab = doc.add_table(rows=1, cols=len(cabecalhos))
    tab.style = "Table Grid"
    tab.autofit = False

    cab = tab.rows[0]
    repetir_cabecalho(cab)
    for i, txt in enumerate(cabecalhos):
        cab.cells[i].text = ""
        cab.cells[i].paragraphs[0].add_run(txt).bold = True

    for valores in linhas:
        l = tab.add_row()
        nao_partir(l)
        for i, v in enumerate(valores):
            forte = v.startswith("**")
            if forte:
                v = v[2:]
            if forte or (i == 0 and negrito_primeira):
                l.cells[i].paragraphs[0].add_run(v).bold = True
            else:
                l.cells[i].text = v

    for l in tab.rows:
        for i, w in enumerate(larguras):
            l.cells[i].width = Cm(w)
    return tab


def nota(doc, texto):
    """Linha pequena e italica, para citar requisitos ou arquivos."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(texto)
    r.font.size = Pt(8.5)
    r.italic = True
    return p


# ===========================================================================
# Saida
# ===========================================================================
def converter_para_pdf(caminho_docx):
    """
    Gera o PDF ao lado do .docx.

    Duas rotas, nesta ordem, porque nem toda maquina tem as duas:

      1. Word, se instalado. E o que produz o resultado fiel.
      2. LibreOffice. Pagina de forma ligeiramente diferente, mas resolve
         numa maquina sem Word.

    Sem nenhum dos dois, avisa e devolve None em vez de derrubar: o .docx, que
    e a entrega, ja esta gravado a essa altura.
    """
    absoluto = os.path.abspath(caminho_docx)
    destino = absoluto[:-5] + ".pdf"

    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            documento = word.Documents.Open(absoluto, False, True)
            documento.SaveAs(destino, FileFormat=17)   # 17 = PDF
            documento.Close(False)
        finally:
            word.Quit()
        return destino
    except Exception as erro:
        print("  Word nao converteu (%s); tentando o LibreOffice." % erro)

    for candidato in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "soffice",
    ):
        try:
            subprocess.run(
                [candidato, "--headless", "--convert-to", "pdf",
                 "--outdir", os.path.dirname(absoluto), absoluto],
                check=True, capture_output=True, timeout=180)
            if os.path.exists(destino):
                return destino
        except (OSError, subprocess.SubprocessError):
            continue

    print("  Nenhum conversor disponivel; so o .docx foi gerado.")
    return None


def salvar(doc, prefixo, data):
    """Grava docs/<prefixo>_AAAA-MM-DD.docx e o PDF ao lado."""
    nome = "docs/%s_%s.docx" % (prefixo, data.isoformat())
    doc.save(nome)
    return nome, converter_para_pdf(nome)


def raiz_do_repositorio():
    """Roda a partir da raiz, e nao de onde o usuario chamou o script."""
    os.chdir(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


def data_do_argumento(argv):
    return datetime.date.fromisoformat(argv[1]) if len(argv) > 1 else datetime.date.today()
